"""
DOCX parser using python-docx with trio support.
"""
import logging
from typing import List, Tuple, Optional, Dict, Any
from pathlib import Path

import trio
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from .base_parser import BaseParser, DocumentStructure, ParsingError

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """
    DOCX parser using python-docx library.
    """
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
        self.parser_name = "DOCXParser"
    
    async def parse(self, file_path: str, **kwargs) -> DocumentStructure:
        """
        Parse DOCX document.
        
        Args:
            file_path: Path to DOCX file
            **kwargs: Additional options:
                - extract_tables: bool
                - preserve_formatting: bool
                
        Returns:
            DocumentStructure with parsed content
        """
        try:
            # Validate file
            if not await self.validate_file(file_path):
                raise ParsingError(f"Invalid DOCX file: {file_path}", self.parser_name, file_path)
            
            # Parse options
            extract_tables = kwargs.get('extract_tables', True)
            preserve_formatting = kwargs.get('preserve_formatting', False)
            
            logger.info(f"Parsing DOCX {file_path}")
            
            # Parse DOCX using python-docx (run in thread)
            doc_structure = await trio.to_thread.run_sync(
                self._parse_docx_sync,
                file_path,
                extract_tables,
                preserve_formatting
            )
            
            logger.info(f"Successfully parsed DOCX {file_path}: {len(doc_structure.text_blocks)} text blocks")
            
            return doc_structure
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise ParsingError(f"DOCX parsing failed: {str(e)}", self.parser_name, file_path)
    
    def _parse_docx_sync(self, file_path: str, extract_tables: bool, preserve_formatting: bool) -> DocumentStructure:
        """
        Synchronous DOCX parsing using python-docx.
        """
        doc_structure = DocumentStructure()
        doc_structure.metadata = self._extract_metadata(file_path)
        
        # Open document
        doc = Document(file_path)
        
        # Process document elements in order
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()
                
                if text:
                    metadata = {
                        "element_type": "paragraph",
                        "style": paragraph.style.name if paragraph.style else None
                    }
                    
                    # Add formatting info if requested
                    if preserve_formatting:
                        metadata["formatting"] = self._extract_paragraph_formatting(paragraph)
                    
                    doc_structure.text_blocks.append((text, metadata))
            
            elif element.tag.endswith('tbl') and extract_tables:  # Table
                table = Table(element, doc)
                table_data = self._extract_table_data(table)
                
                if table_data["content"]:
                    doc_structure.tables.append(table_data)
                    
                    # Also add table as text block
                    metadata = {
                        "element_type": "table",
                        "rows": table_data["rows"],
                        "columns": table_data["columns"]
                    }
                    doc_structure.text_blocks.append((table_data["content"], metadata))
        
        # Update metadata
        doc_structure.metadata.update({
            "total_paragraphs": len([b for b in doc_structure.text_blocks if b[1].get("element_type") == "paragraph"]),
            "total_tables": len(doc_structure.tables),
            "total_text_blocks": len(doc_structure.text_blocks)
        })
        
        return doc_structure
    
    def _extract_table_data(self, table: Table) -> Dict[str, Any]:
        """
        Extract data from DOCX table.
        """
        rows_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            rows_data.append(row_data)
        
        # Convert to text representation
        table_text = []
        for row in rows_data:
            table_text.append(" | ".join(row))
        
        return {
            "content": "\n".join(table_text),
            "rows": len(rows_data),
            "columns": len(rows_data[0]) if rows_data else 0,
            "raw_data": rows_data,
            "element_type": "table"
        }
    
    def _extract_paragraph_formatting(self, paragraph: Paragraph) -> Dict[str, Any]:
        """
        Extract formatting information from paragraph.
        """
        formatting = {}
        
        try:
            # Paragraph formatting
            if paragraph.paragraph_format:
                pf = paragraph.paragraph_format
                formatting.update({
                    "alignment": str(pf.alignment) if pf.alignment else None,
                    "left_indent": pf.left_indent,
                    "right_indent": pf.right_indent,
                    "first_line_indent": pf.first_line_indent,
                    "space_before": pf.space_before,
                    "space_after": pf.space_after
                })
            
            # Run formatting (for first run)
            if paragraph.runs:
                run = paragraph.runs[0]
                if run.font:
                    formatting.update({
                        "font_name": run.font.name,
                        "font_size": run.font.size,
                        "bold": run.font.bold,
                        "italic": run.font.italic,
                        "underline": run.font.underline
                    })
        
        except Exception as e:
            logger.debug(f"Failed to extract formatting: {e}")
        
        return formatting